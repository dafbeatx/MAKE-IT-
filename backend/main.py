import os
import io
import base64
import tempfile
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from engine import apply_margin, apply_typography, apply_heading, format_heading_text

load_dotenv()

app = FastAPI(
    title="MAKE!T Document Generator API",
    version="2.0.0",
    description="Template-based document generator. No AI — users write content, system formats it.",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from pydantic import BaseModel, Field

# ─── Models ─────────────────────────────────────────────

class Section(BaseModel):
    id: str = Field(..., min_length=1)
    title: str = Field(..., min_length=1, description="Section heading title")
    content: str
    quran: Optional[str] = None
    footnote: Optional[str] = None

class Chapter(BaseModel):
    id: str = Field(..., min_length=1)
    title: str = Field(..., min_length=1, description="Chapter heading title")
    sections: List[Section]

class Identity(BaseModel):
    title: str = Field(..., min_length=1, description="Document main title")
    docSubtype: Optional[str] = None
    name: str = Field(..., min_length=1)
    nim: str = Field(..., min_length=1)
    institution: str = Field(..., min_length=1)
    faculty: str = Field(..., min_length=1)
    prodi: Optional[str] = None
    supervisor: str
    year: str = Field(..., min_length=1)
    year_hijri: Optional[str] = None
    degree_purpose: Optional[str] = None
    logo: Optional[str] = None
    address: Optional[str] = None

class FormatConfig(BaseModel):
    font_name: str = "Times New Roman"
    font_size_body: int = 12
    font_size_heading: int = 14       # legacy alias for h1
    font_size_h1: Optional[int] = None
    font_size_h2: Optional[int] = None
    font_size_h3: Optional[int] = None
    font_size_subtitle: Optional[int] = None
    line_spacing: float = 1.5
    h1_bold: bool = True
    h1_uppercase: bool = True
    h1_center: bool = True
    h2_bold: bool = True
    h3_bold: bool = True
    margin_top: float = 4.0
    margin_bottom: float = 3.0
    margin_left: float = 4.0
    margin_right: float = 3.0
    page_number_prelim: str = "roman"     # "roman" | "arabic" | "none"
    page_number_body: str = "arabic"      # "arabic" | "roman"
    numbering_system: str = "standard-indo"  # "standard-indo" | "numeric" | "custom"
    first_line_indent: float = 1.27
    has_quran: bool = False
    has_footnote: bool = False
    has_abstract: bool = False

    @property
    def h1_size(self) -> int:
        return self.font_size_h1 or self.font_size_heading

    @property
    def h2_size(self) -> int:
        return self.font_size_h2 or self.font_size_body

    @property
    def h3_size(self) -> int:
        return self.font_size_h3 or self.font_size_body

class GenerateRequest(BaseModel):
    docType: Optional[str] = None
    has_cover: Optional[bool] = False
    identity: Identity
    chapters: List[Chapter]
    format_config: FormatConfig = FormatConfig()
    abstract_paragraphs: Optional[List[str]] = None

# ─── Format Presets ─────────────────────────────────────

FORMAT_PRESETS = {
    "standar-a": FormatConfig(
        font_name="Times New Roman", font_size_body=12, font_size_heading=14,
        line_spacing=1.5, margin_top=4.0, margin_bottom=3.0, margin_left=4.0, margin_right=3.0,
    ),
    "standar-b": FormatConfig(
        font_name="Arial", font_size_body=11, font_size_heading=12,
        line_spacing=1.15, margin_top=2.54, margin_bottom=2.54, margin_left=2.54, margin_right=2.54,
    ),
    "standar-c": FormatConfig(
        font_name="Calibri", font_size_body=11, font_size_heading=13,
        line_spacing=1.5, margin_top=2.5, margin_bottom=2.5, margin_left=3.0, margin_right=2.5,
    ),
}

# ─── Document Generation Logic ─────────────────────────

def build_document(req: GenerateRequest) -> Document:
    doc = Document()
    fmt = req.format_config

    # Set margins using Rule Engine
    for section in doc.sections:
        apply_margin(section, fmt)

    # ── Cover Page ──
    # Only generated when user explicitly enables "Sampul Depan" in Step 2.
    # No placeholder text — only renders data the user has actually filled in.
    if req.has_cover:

        from markdown_parser import parse_markdown_to_runs
        def _add_centered_line(text: str, font_size: int, bold: bool = False, italic: bool = False, spacing_after: int = 0):
            """Helper to add a single centered line with consistent formatting, parsing Markdown tokens using external parser."""
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if spacing_after:
                p.paragraph_format.space_after = Pt(spacing_after)
            
            parse_markdown_to_runs(p, text, apply_typography, fmt, font_size, base_bold=bold)

        def _add_logo(width_inches: float):
            """Helper to safely decode and insert a base64 logo."""
            if not req.identity.logo:
                return
            try:
                _header, encoded = req.identity.logo.split(",", 1)
                image_data = base64.b64decode(encoded)
                logo_io = io.BytesIO(image_data)
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_io, width=Inches(width_inches))
            except Exception as e:
                print("Failed to decode logo:", e)

        # 1. Judul Skripsi (from user data)
        if req.identity.title:
            title_text = req.identity.title
            if fmt.h1_uppercase:
                title_text = title_text.upper()
                
            _add_centered_line(
                title_text,
                fmt.h1_size,
                bold=fmt.h1_bold,
                spacing_after=12,
            )

        # 2. Doc Subtype (SKRIPSI / SEMINAR HASIL / etc.)
        if req.identity.docSubtype:
            _add_centered_line(
                req.identity.docSubtype.upper(),
                fmt.h2_size,
                bold=fmt.h2_bold,
                spacing_after=12,
            )

        doc.add_paragraph("")

        # 3. Degree Purpose Text
        if req.identity.degree_purpose:
            _add_centered_line(
                req.identity.degree_purpose,
                fmt.font_size_body,
                italic=True,
            )

        for _ in range(2):
            doc.add_paragraph("")

        # 4. Middle Logo (single, large, centered - 5x5cm = 1.97 inches)
        _add_logo(1.97)

        for _ in range(1):
            doc.add_paragraph("")

        # 5. "Disusun Oleh:"
        _add_centered_line("Disusun oleh :", fmt.font_size_body)

        # 6. Name & NIM
        if req.identity.name:
            _add_centered_line(req.identity.name, fmt.font_size_body, bold=False)

        if req.identity.nim:
            _add_centered_line(req.identity.nim, fmt.font_size_body, bold=False)

        for _ in range(4):
            doc.add_paragraph("")

        # 7. Institution block (Prodi → Fakultas → Institusi)
        institution_lines: list[str] = []
        if req.identity.prodi:
            institution_lines.append(req.identity.prodi.upper())
        if req.identity.faculty:
            institution_lines.append(req.identity.faculty.upper())
        if req.identity.institution:
            institution_lines.append(req.identity.institution.upper())

        for text in institution_lines:
            _add_centered_line(text, fmt.font_size_body, bold=True)

        # 8. Address
        if req.identity.address:
            _add_centered_line(req.identity.address.upper(), fmt.font_size_body, bold=True)

        # 9. Years (Masehi / Hijriah)
        year_parts = [p for p in [req.identity.year, req.identity.year_hijri] if p]
        if year_parts:
            _add_centered_line(" / ".join(year_parts), fmt.font_size_body, bold=True)

        doc.add_page_break()

    # ── Abstract ──
    from markdown_parser import parse_markdown_to_runs

    if fmt.has_abstract and req.abstract_paragraphs:
        doc.add_page_break()
        abs_h = doc.add_paragraph()
        apply_heading(abs_h, 1, fmt)
        parse_markdown_to_runs(abs_h, "ABSTRAK", apply_typography, fmt, fmt.h1_size, base_bold=fmt.h1_bold)
        doc.add_paragraph("")
        
        for para_text in req.abstract_paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(fmt.first_line_indent)
                p.paragraph_format.line_spacing = fmt.line_spacing
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parse_markdown_to_runs(p, para_text.strip(), apply_typography, fmt, fmt.font_size_body)

        doc.add_page_break()

    # ── Chapters ──
    for chapter in req.chapters:
        chapter_title_parts = chapter.title.split(":", 1)
        bab_number = chapter_title_parts[0].strip()
        bab_name = chapter_title_parts[1].strip() if len(chapter_title_parts) > 1 else ""

        bab_number = format_heading_text(bab_number, 1, fmt)
        bab_name = format_heading_text(bab_name, 1, fmt)

        h_para = doc.add_paragraph()
        apply_heading(h_para, 1, fmt)
        parse_markdown_to_runs(h_para, bab_number, apply_typography, fmt, fmt.h1_size, base_bold=fmt.h1_bold)

        if bab_name:
            h_para2 = doc.add_paragraph()
            apply_heading(h_para2, 1, fmt)
            parse_markdown_to_runs(h_para2, bab_name, apply_typography, fmt, fmt.h1_size, base_bold=fmt.h1_bold)

        doc.add_paragraph("")

        for section in chapter.sections:
            s_para = doc.add_paragraph()
            apply_heading(s_para, 2, fmt)
            parse_markdown_to_runs(s_para, section.title, apply_typography, fmt, fmt.h2_size, base_bold=fmt.h2_bold)

            # Content paragraphs
            if section.content.strip():
                for para_text in section.content.strip().split("\n"):
                    if para_text.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(fmt.first_line_indent)
                        p.paragraph_format.line_spacing = fmt.line_spacing
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        parse_markdown_to_runs(p, para_text.strip(), apply_typography, fmt, fmt.font_size_body)

            # Insert Quran text if configured and present
            if fmt.has_quran and section.quran and section.quran.strip():
                q_para = doc.add_paragraph()
                q_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                q_para.paragraph_format.line_spacing = 2.0
                q_run = q_para.add_run(section.quran.strip())
                # Quran is typically rendered larger
                q_run.font.size = Pt(fmt.font_size_body + 6)
                
            # Insert Endnotes/Footnotes workaround if configured and present
            if fmt.has_footnote and section.footnote and section.footnote.strip():
                fn_para = doc.add_paragraph()
                fn_para.paragraph_format.space_before = Pt(6)
                fn_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fn_run = fn_para.add_run("______________________________" + "\n" + section.footnote.strip())
                fn_run.font.name = fmt.font_name
                fn_run.font.size = Pt(10)

            doc.add_paragraph("")

        doc.add_page_break()

    return doc


# ─── Endpoints ──────────────────────────────────────────

@app.get("/")
def read_root():
    return {"app": "MAKE!T", "version": "2.0.0", "ai": False}


@app.post("/api/generate")
async def generate_document(req: GenerateRequest):
    """Generate a .docx from user-provided content. No AI involved."""
    try:
        doc = build_document(req)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"MAKEIT_{req.identity.name.replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/rapikan")
async def rapikan_document(
    file: UploadFile = File(...),
    format_preset: str = Form("standar-a"),
):
    """Upload a .docx and apply academic formatting without changing text content."""
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="File harus berformat .docx")

    fmt = FORMAT_PRESETS.get(format_preset, FORMAT_PRESETS["standar-a"])

    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))

        # Apply margins
        for section in doc.sections:
            section.top_margin = Cm(fmt.margin_top)
            section.bottom_margin = Cm(fmt.margin_bottom)
            section.left_margin = Cm(fmt.margin_left)
            section.right_margin = Cm(fmt.margin_right)

        # Apply font and spacing to all paragraphs
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = fmt.line_spacing
            for run in para.runs:
                run.font.name = fmt.font_name
                run.font.size = Pt(fmt.font_size_body)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="RAPIKAN_{file.filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/spss-table")
async def process_spss(file: UploadFile = File(...)):
    """Upload .sav or .xlsx → generate descriptive statistics table in Word."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("sav", "xlsx", "xls"):
        raise HTTPException(status_code=400, detail="Format harus .sav, .xlsx, atau .xls")

    try:
        content = await file.read()
        tmp_path = tempfile.mktemp(suffix=f".{ext}")
        with open(tmp_path, "wb") as f:
            f.write(content)

        if ext == "sav":
            import pyreadstat
            df, meta = pyreadstat.read_sav(tmp_path)
        else:
            import pandas as pd
            df = pd.read_excel(tmp_path)

        # Build Word doc with descriptive stats table
        doc = Document()
        
        # Helper to apply clean styles
        def format_table(table):
            table.style = "Light Grid Accent 1"
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(10)

        # 1. Descriptive Statistics
        doc.add_heading("Tabel Statistik Deskriptif", level=2)
        stats = df.describe().round(3)
        table = doc.add_table(rows=1, cols=len(stats.columns) + 1)
        
        hdr = table.rows[0].cells
        hdr[0].text = "Statistik"
        for i, col in enumerate(stats.columns):
            hdr[i + 1].text = str(col)

        for idx, row in stats.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(idx)
            for i, val in enumerate(row):
                cells[i + 1].text = str(val)
        format_table(table)
        doc.add_paragraph()

        # 2. Normality Test (Shapiro-Wilk)
        from scipy import stats as scipy_stats
        import numpy as np

        doc.add_heading("Uji Normalitas (Shapiro-Wilk)", level=2)
        norm_table = doc.add_table(rows=1, cols=3)
        n_hdr = norm_table.rows[0].cells
        n_hdr[0].text = "Variabel"
        n_hdr[1].text = "Statistic (W)"
        n_hdr[2].text = "p-value"

        numeric_cols = df.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            clean_data = df[col].dropna()
            if len(clean_data) >= 3:
                stat, p_val = scipy_stats.shapiro(clean_data)
                cells = norm_table.add_row().cells
                cells[0].text = str(col)
                cells[1].text = f"{stat:.4f}"
                cells[2].text = f"{p_val:.4f} {'(Normal)' if p_val > 0.05 else '(Tidak Normal)'}"
        format_table(norm_table)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)


        os.unlink(tmp_path)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="Tabel_{file.filename}.docx"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/template/render")
async def render_custom_template(
    template: UploadFile = File(...),
    data: str = Form("{}"),
):
    """
    Upload a custom .docx template with Jinja2 {{ placeholders }}.
    Pass context data as JSON string to fill in the placeholders.
    """
    import json
    from template_renderer import render_uploaded_template  # type: ignore

    if not template.filename or not template.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Template harus .docx")

    try:
        context = json.loads(data)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Data harus berformat JSON valid")

    try:
        buffer = await render_uploaded_template(template, context)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="rendered_{template.filename}"'
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
